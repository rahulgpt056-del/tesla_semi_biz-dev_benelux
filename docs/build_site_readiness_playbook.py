"""
Generate the Tesla Semi Benelux Depot & Grid Site-Readiness Playbook (.docx).

Run:    python docs/build_site_readiness_playbook.py
Output: docs/Tesla-Semi-Benelux-Site-Readiness-Playbook.docx
"""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

ROOT = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(ROOT, "Tesla-Semi-Benelux-Site-Readiness-Playbook.docx")

NAVY = RGBColor(0x0B, 0x25, 0x45)
RED = RGBColor(0xE3, 0x19, 0x37)
GREY = RGBColor(0x5B, 0x66, 0x70)

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = NAVY
    return p


def body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    doc.add_paragraph(text, style="List Number")


def table_section(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = htext
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = NAVY
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


# --------------------------------------------------------------------------- #
# Title page
# --------------------------------------------------------------------------- #
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Tesla Semi · Benelux")
run.font.size = Pt(30)
run.font.color.rgb = NAVY
run.bold = True

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Depot & Grid Site-Readiness Playbook")
run.font.size = Pt(18)
run.font.color.rgb = RED
run.bold = True

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run(
    "A working checklist for taking a Benelux fleet account from first "
    "conversation to a commissioned depot charging site.\n"
    "Prepared for: Business Development Manager, Semi — Benelux (Req. ID 267880)"
)
run.font.size = Pt(11)
run.font.color.rgb = GREY

doc.add_page_break()

# --------------------------------------------------------------------------- #
# Purpose
# --------------------------------------------------------------------------- #
h1("Why this playbook exists")
body(
    "The slowest part of a Semi deal is rarely the vehicle decision — it's "
    "getting a depot from 'interested' to 'grid-connected and permitted for "
    "Megachargers'. Utility connection requests in the Netherlands and "
    "Belgium routinely take 6-18 months when capacity upgrades are needed. "
    "This playbook front-loads those questions so site-readiness work starts "
    "in parallel with the commercial conversation, not after it."
)
body(
    "It is organised around six stages that mirror the deal lifecycle in the "
    "Tesla Semi BD role: qualify the fleet, spec the vehicle and charging "
    "solution, confirm utility readiness, progress the Sales Contract, plan "
    "onboarding, and commission the site."
)

# --------------------------------------------------------------------------- #
# Stage 1
# --------------------------------------------------------------------------- #
h1("Stage 1 — Qualify the fleet")
body(
    "Goal: confirm the account's duty cycle is a realistic Semi fit before "
    "investing further time. Use the TCO & Duty-Cycle Calculator "
    "(app/streamlit_app.py) with real or estimated figures from the account."
)
table_section(
    ["Question", "Why it matters", "Where it's used"],
    [
        ("What is the daily distance per truck, and is it return-to-depot?", "Determines if 800km Semi range covers the route without mid-route charging.", "TCO calculator — Fleet & duty cycle inputs"),
        ("What is the typical payload?", "Heavier payloads increase energy consumption above the 19t reference point.", "TCO calculator — effective consumption"),
        ("How many hours does each truck dwell at the depot overnight?", "Determines whether overnight charging at a realistic charger power can fully replenish the battery.", "TCO calculator — charge time vs. dwell"),
        ("What is the account's current grid connection capacity at the depot?", "First filter for whether a grid upgrade will be needed at all.", "TCO calculator — depot capacity vs. peak load"),
        ("Does the account have a public sustainability commitment or EV pilot?", "Signals budget availability and internal champion for a ZE pilot.", "Territory map — sustainability_signal"),
    ],
)
body("Outcome of this stage: a duty-cycle fit verdict — Fit, Conditional, or No fit — for at least one representative route.", bold=True)

# --------------------------------------------------------------------------- #
# Stage 2
# --------------------------------------------------------------------------- #
h1("Stage 2 — Spec the vehicle and charging solution")
body("Goal: turn the qualified duty cycle into a concrete fleet size, charger count, and depot layout.")
bullet("Confirm the number of trucks in the initial pilot vs. the account's full addressable fleet (use the territory map's fleet_scale as a starting estimate).")
bullet("Use the TCO calculator's 'Site & charging' tab to determine the number of chargers needed and the resulting peak load (kW).")
bullet("Compare peak load against the depot's usable grid capacity (90% of nameplate, per the calculator's UsableCap assumption).")
bullet("If peak load exceeds usable capacity, flag a grid upgrade requirement immediately — this becomes the critical path item for Stage 3.")
bullet("Decide on charger power tier: depot chargers (lower kW, slower overnight fill) vs. Megacharger (up to 1.2 MW, faster but higher grid demand).")

# --------------------------------------------------------------------------- #
# Stage 3
# --------------------------------------------------------------------------- #
h1("Stage 3 — Confirm utility readiness")
body(
    "Goal: get the grid-connection question into the relevant Distribution "
    "System Operator's (DSO) queue as early as possible, since this is "
    "almost always the longest lead-time item in the project."
)
table_section(
    ["Country / region", "Relevant DSO(s)", "Typical first step"],
    [
        ("Netherlands — most regions", "Liander, Enexis, Stedin (depends on depot location)", "Submit a capacity-increase request (\"aansluiting vergroten\") via the DSO's business portal; request an indicative lead time and cost estimate."),
        ("Belgium — Flanders", "Fluvius", "Submit a connection study request; ask about available headroom on the local MV feeder near the depot."),
        ("Belgium — Wallonia", "ORES / Resa (depends on province)", "Same as above via the relevant regional DSO."),
        ("Luxembourg", "Creos Luxembourg", "Contact Creos for a feasibility check on MV connection capacity at the site."),
    ],
)
bullet("If a capacity upgrade is needed, request the DSO's indicative timeline in writing at first contact — this becomes the master schedule's critical path.")
bullet("In parallel, check whether the depot's existing connection can support a phased rollout (e.g. first 2-3 chargers within existing headroom while the upgrade is processed).")
bullet("Confirm whether the account already has an on-site substation or transformer that could be upgraded rather than requiring a new grid connection point.")

# --------------------------------------------------------------------------- #
# Stage 4
# --------------------------------------------------------------------------- #
h1("Stage 4 — Sales Contract & incentive stacking")
body("Goal: assemble the commercial package so the Sales Contract reflects the true net cost to the customer.")
bullet("Confirm AanZET (Aanschafsubsidie Zero-Emissie Trucks) eligibility and current budget-round status with RVO — subsidy amounts and availability change per round.")
bullet("Confirm the vehicle's vrachtwagenheffing (Dutch truck toll) status — zero-emission trucks are expected to receive a reduced/zero tariff; verify the current rule before quoting savings.")
bullet("Use the TCO calculator's 5-year saving and payback outputs as the basis for the commercial narrative in the proposal.")
bullet("Identify whether the account qualifies for any regional or EU co-funding programmes for charging infrastructure (in addition to the vehicle subsidy).")
bullet("Document any site-specific capex (grid upgrade cost, civil works for charger pads) as a line item — this is often the largest non-vehicle cost.")

# --------------------------------------------------------------------------- #
# Stage 5
# --------------------------------------------------------------------------- #
h1("Stage 5 — Onboarding plan")
body("Goal: prepare the account's operations team for the transition, not just the finance team.")
bullet("Driver familiarisation: range expectations, regenerative braking, charging routine.")
bullet("Depot operations: charging schedule design so trucks queue efficiently within the dwell window — revisit the TCO calculator's load profile to confirm the charging window doesn't create a new peak-demand charge.")
bullet("Maintenance: confirm service network coverage for the depot's region and agree SLAs.")
bullet("Telematics/fleet management integration: confirm how Semi data feeds into the account's existing fleet management system.")

# --------------------------------------------------------------------------- #
# Stage 6
# --------------------------------------------------------------------------- #
h1("Stage 6 — Commissioning")
body("Goal: go live with a working, monitored charging site and a feedback loop back into the BD relationship.")
bullet("Joint commissioning checklist with the DSO, the charging-infrastructure installer, and the account's facilities team.")
bullet("Confirm metering and billing setup for the new connection (especially if a separate sub-meter is used for EV charging tariffs).")
bullet("Schedule a 30/60/90-day post-go-live review to capture real-world energy use vs. the TCO calculator's projections — feed any material gaps back into the assumptions used for the next account.")
bullet("Use a successful commissioning as a reference case for the next Tier A/B account in the territory map.")

# --------------------------------------------------------------------------- #
# Risk register
# --------------------------------------------------------------------------- #
h1("Risk register")
table_section(
    ["Risk", "Likelihood", "Mitigation"],
    [
        ("DSO grid-upgrade lead time exceeds the account's vehicle delivery timeline", "High", "Submit the connection request at Stage 1, in parallel with commercial qualification — not after contract signature."),
        ("AanZET budget round closes or is oversubscribed before contract signature", "Medium", "Track RVO budget-round status from first qualified conversation; build urgency into the proposal timeline."),
        ("Real-world energy consumption exceeds the model's route-factor assumption", "Medium", "Use a short telematics-based pilot on the representative route before committing to fleet-wide charger sizing."),
        ("Vrachtwagenheffing zero-emission exemption rules change before go-live", "Low-Medium", "Keep the TCO calculator's toll assumption flagged as 'confirm before contracting' and re-check at Stage 4."),
        ("Depot dwell time is shorter than assumed (e.g. multi-shift operations)", "Medium", "Re-run the TCO calculator's charge-time-vs-dwell check with actual shift patterns before finalising charger count."),
    ],
)

# --------------------------------------------------------------------------- #
# Appendix
# --------------------------------------------------------------------------- #
h1("Appendix — supporting tools")
bullet("TCO & Duty-Cycle Calculator (app/streamlit_app.py) — fit verdict, energy/charging plan, 5-year TCO and payback for any account.")
bullet("Target-Account Territory Map (app/territory_app.py) — tiered Benelux fleet-operator list to sequence outreach.")
bullet("Both tools share a single assumptions file (data/assumptions.json) that refreshes weekly via GitHub Actions, so the numbers in this playbook stay current.")

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("WROTE", OUT)
